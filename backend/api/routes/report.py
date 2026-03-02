"""
Report Route — Generates and serves professional DOCX compliance reports.
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json
import aiosqlite
from datetime import datetime

router = APIRouter()
DB_PATH = "reg_validator.db"


def _style_cell(cell, text, bold=False, size=9, color=None):
    """Style a table cell with font settings."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _add_styled_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)  # slate-800
    return heading


@router.get("/{session_id}/download")
async def download_report(session_id: str):
    """Generate and download a professional compliance report."""
    async with aiosqlite.connect(DB_PATH) as db:
        db.row_factory = aiosqlite.Row
        async with db.execute(
            "SELECT * FROM validation_results WHERE session_id = ?", (session_id,)
        ) as cursor:
            row = await cursor.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Results not found")

            data = json.loads(row["result_json"])
            accuracy_score = row["accuracy_score"]
            consistency_score = row["consistency_score"]

            # Load pipeline metadata if available
            pipeline_path = os.path.join("storage/sessions", session_id, "pipeline_output.json")
            corrections = []
            accuracy_breakdown = {}
            if os.path.exists(pipeline_path):
                with open(pipeline_path, "r") as f:
                    pipeline_data = json.load(f)
                    corrections = pipeline_data.get("corrections", [])
                    accuracy_breakdown = pipeline_data.get("accuracy", {})

            # ── Build DOCX ───────────────────────────────────────────
            doc = Document()

            # Title
            title = doc.add_heading("Regulatory Compliance\nAudit Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

            # Metadata section
            doc.add_paragraph("")  # spacer
            meta_table = doc.add_table(rows=5, cols=2)
            meta_table.style = "Light Grid Accent 1"
            meta_data = [
                ("Document", row["document_name"]),
                ("Guideline", row["guideline_name"]),
                ("Session ID", session_id),
                ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M UTC")),
                ("Pipeline Version", "2.0"),
            ]
            for i, (label, value) in enumerate(meta_data):
                _style_cell(meta_table.rows[i].cells[0], label, bold=True, size=10)
                _style_cell(meta_table.rows[i].cells[1], value, size=10)

            # ── Scores Summary ────────────────────────────────────────
            _add_styled_heading(doc, "Compliance Summary", level=1)
            
            score_table = doc.add_table(rows=1, cols=3)
            score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = ["Compliance Score", "Accuracy Score", "Consistency Score"]
            values = [
                f"{data.get('compliance_score', 'N/A')}%",
                f"{accuracy_score}%",
                f"{consistency_score}%"
            ]
            for i, (h, v) in enumerate(zip(headers, values)):
                _style_cell(score_table.rows[0].cells[i], f"{h}\n{v}", bold=True, size=11)
                score_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Risk summary
            risk = data.get("risk_summary", {})
            if risk:
                doc.add_paragraph("")
                risk_text = (
                    f"Risk Distribution: {risk.get('critical', 0)} Critical | "
                    f"{risk.get('major', 0)} Major | {risk.get('minor', 0)} Minor | "
                    f"{risk.get('observations', 0)} Observations"
                )
                p = doc.add_paragraph(risk_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

            # ── Executive Summary ─────────────────────────────────────
            _add_styled_heading(doc, "Executive Summary", level=1)
            summary_text = data.get("executive_summary", "No executive summary was generated.")
            p = doc.add_paragraph(summary_text)
            for run in p.runs:
                run.font.size = Pt(11)

            # ── Detailed Findings ─────────────────────────────────────
            _add_styled_heading(doc, "Detailed Findings", level=1)

            findings = data.get("findings", [])
            total = len(findings)
            doc.add_paragraph(f"Total clauses reviewed: {total}")

            if findings:
                table = doc.add_table(rows=1, cols=5)
                table.style = "Light Grid Accent 1"
                table.autofit = True

                # Headers
                header_cells = table.rows[0].cells
                for i, header in enumerate(["Clause", "Status", "Severity", "Description", "Recommendation"]):
                    _style_cell(header_cells[i], header, bold=True, size=9, 
                              color=RGBColor(0xff, 0xff, 0xff))

                # Status color mapping
                status_colors = {
                    "Compliant": RGBColor(0x10, 0xb9, 0x81),
                    "Non-Compliant": RGBColor(0xef, 0x44, 0x44),
                    "Partial": RGBColor(0xf5, 0x9e, 0x0b),
                    "Observation": RGBColor(0x3b, 0x82, 0xf6),
                    "Not Applicable": RGBColor(0x94, 0xa3, 0xb8),
                }

                for finding in findings:
                    row_cells = table.add_row().cells
                    _style_cell(row_cells[0], finding.get("clause", "—"), size=8)
                    
                    status = finding.get("status", "—")
                    _style_cell(row_cells[1], status, bold=True, size=8,
                              color=status_colors.get(status, RGBColor(0, 0, 0)))
                    
                    _style_cell(row_cells[2], finding.get("severity", "N/A"), size=8)
                    _style_cell(row_cells[3], finding.get("description", "—"), size=8)
                    _style_cell(row_cells[4], finding.get("recommendation", "—"), size=8)

            # ── Corrections & Audit Trail ─────────────────────────────
            if corrections:
                _add_styled_heading(doc, "Pipeline Corrections Applied", level=2)
                for c in corrections:
                    doc.add_paragraph(f"• {c}", style="List Bullet")

            # ── Accuracy Breakdown ────────────────────────────────────
            breakdown = accuracy_breakdown.get("breakdown", {})
            if breakdown:
                _add_styled_heading(doc, "Accuracy Breakdown", level=2)
                for metric, score in breakdown.items():
                    label = metric.replace("_", " ").title()
                    doc.add_paragraph(f"{label}: {score}%")

            # ── Footer / Disclaimer ──────────────────────────────────
            doc.add_paragraph("")
            disclaimer = doc.add_paragraph(
                "DISCLAIMER: This report was generated by an AI-assisted regulatory validation engine. "
                "It is intended as a decision-support tool and does not constitute formal regulatory advice. "
                "All findings should be reviewed by qualified regulatory affairs professionals before submission."
            )
            for run in disclaimer.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
                run.italic = True

            # Save
            report_path = os.path.join("storage/sessions", session_id, "report.docx")
            doc.save(report_path)

            return FileResponse(
                path=report_path,
                filename=f"ComplianceLens_Report_{session_id[:8]}.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
